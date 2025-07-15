import openai
import googlemaps
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import json
import requests
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
import logging
from pathlib import Path
import re
from geopy.geocoders import GoogleV3
from geopy.exc import GeocoderTimedOut
import time
from io import BytesIO

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class PropertyReportData:
    """Complete data structure for property report"""
    # Basic Info
    address: str
    date: str = field(default_factory=lambda: datetime.now().strftime("%B %d, %Y"))
    
    # Preparer Info
    prepared_by: str = ""
    prepared_by_company: str = ""
    prepared_by_address: str = ""
    
    # Client Info  
    prepared_for: str = ""
    prepared_for_company: str = ""
    prepared_for_address: str = ""
    
    # Property Details
    property_name: str = ""
    property_type: str = ""
    state: str = ""
    county: str = ""
    longitude: str = ""
    latitude: str = ""
    
    # Physical Characteristics (mostly static)
    topography: str = "Level at street Grade"
    shape: str = "Irregular"
    access: str = "Good"  # Average, Average/Good, Good, Good/Excellent, Excellent
    exposure: str = "Average/Good"  # Average, Average/Good, Good, Good/Excellent, Excellent
    
    # Property Specific
    lot_area: str = ""
    acres: str = ""
    recorded_sale_date: str = ""
    zoning: str = ""
    apn: str = ""
    current_owner: str = ""
    
    # Market Analysis (static)
    marketing_period: str = "Six months or less"
    
    # SWOT Analysis
    swot_strengths: str = ""
    swot_weaknesses: str = ""
    swot_opportunities: str = ""
    swot_threats: str = ""
    
    # Generated Content
    property_summary: str = ""
    location_summary: str = ""
    demographic_analysis: str = ""
    size_and_topography: str = ""
    population_analysis: str = ""
    household_trends: str = ""
    employment_analysis: str = ""
    economic_factors: str = ""
    community_services: str = ""
    
    # Image paths
    aerial_image_path: Optional[str] = None
    street_view_image_path: Optional[str] = None

class ComprehensivePropertyReportGenerator:
    def __init__(self, openai_api_key: str, google_api_key: str, template_path: str, output_dir: str = "output"):
        """
        Initialize the Comprehensive Property Report Generator
        
        Args:
            openai_api_key: OpenAI API key
            google_api_key: Google Maps API key
            template_path: Path to the Word document template
            output_dir: Directory to save generated reports
        """
        self.openai_client = openai.OpenAI(api_key=openai_api_key)
        self.gmaps = googlemaps.Client(key=google_api_key)
        self.google_api_key = google_api_key
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Create images subdirectory
        self.images_dir = self.output_dir / "images"
        self.images_dir.mkdir(exist_ok=True)
        
        # Verify template exists
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")

    def get_property_images(self, address: str, lat: float, lng: float) -> Tuple[Optional[str], Optional[str]]:
        """
        Get aerial and street view images for the property
        
        Args:
            address: The property address
            lat: Latitude
            lng: Longitude
            
        Returns:
            Tuple of (aerial_image_path, street_view_image_path)
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        aerial_path = None
        street_view_path = None
        
        try:
            # 1. Get Aerial/Satellite Image
            logger.info("Fetching aerial image...")
            aerial_params = {
                "center": f"{lat},{lng}",
                "zoom": "18",
                "size": "1000x1000",
                "maptype": "hybrid",
                "key": self.google_api_key,
                "markers": f"{lat},{lng}"
            }
            
            aerial_response = requests.get(
                "https://maps.googleapis.com/maps/api/staticmap", 
                params=aerial_params,
                timeout=30
            )
            aerial_response.raise_for_status()
            
            aerial_filename = f"aerial_{timestamp}.jpg"
            aerial_path = self.images_dir / aerial_filename
            with open(aerial_path, "wb") as f:
                f.write(aerial_response.content)
            logger.info(f"Aerial image saved: {aerial_path}")
            
            # 2. Get Street View Image
            logger.info("Fetching Street View image...")
            street_view_params = {
                "size": "600x500",
                "location": address,
                "pitch": "0",
                "fov": "90",
                "key": self.google_api_key
            }
            
            street_view_response = requests.get(
                "https://maps.googleapis.com/maps/api/streetview", 
                params=street_view_params,
                timeout=30
            )
            street_view_response.raise_for_status()
            
            # Check if Street View is available
            if street_view_response.headers.get('content-type', '').startswith('image/'):
                street_view_filename = f"street_view_{timestamp}.jpg"
                street_view_path = self.images_dir / street_view_filename
                with open(street_view_path, "wb") as f:
                    f.write(street_view_response.content)
                logger.info(f"Street view image saved: {street_view_path}")
            else:
                logger.warning("Street View not available for this address")
                
        except Exception as e:
            logger.error(f"Error fetching images: {e}")
            
        return str(aerial_path) if aerial_path else None, str(street_view_path) if street_view_path else None

    def get_coordinates_and_details(self, address: str) -> Tuple[float, float, Dict]:
        """
        Get latitude, longitude, and detailed information from Google Geocoding API
        """
        try:
            geocode_result = self.gmaps.geocode(address)
            if not geocode_result:
                raise ValueError(f"No geocoding results for address: {address}")
            
            location = geocode_result[0]['geometry']['location']
            lat = location['lat']
            lng = location['lng']
            
            # Extract address components
            components = geocode_result[0]['address_components']
            address_details = {}
            
            for component in components:
                types = component['types']
                if 'administrative_area_level_2' in types:
                    address_details['county'] = component['long_name']
                elif 'administrative_area_level_1' in types:
                    address_details['state'] = component['long_name']
                elif 'locality' in types:
                    address_details['city'] = component['long_name']
                elif 'postal_code' in types:
                    address_details['zip_code'] = component['long_name']
            
            return lat, lng, address_details
            
        except Exception as e:
            logger.error(f"Error getting coordinates for {address}: {e}")
            raise

    def get_census_data(self, lat: float, lng: float, county: str, state: str) -> Dict:
        """
        Get demographic data from US Census API
        """
        try:
            # This is a simplified version - you'd need to implement actual Census API calls
            # For now, we'll generate realistic data through AI
            census_prompt = f"""
            Generate realistic and current demographic data for coordinates {lat}, {lng} in {county}, {state}.
            Include population statistics, household data, employment data, and economic factors.
            Format as JSON with keys: population_2020, population_growth_rate, households_2020, 
            avg_household_size, employment_rate, major_industries, median_income.
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a demographic data analyst. Provide realistic census-like data in JSON format."},
                    {"role": "user", "content": census_prompt}
                ],
                temperature=0.1
            )
            
            return json.loads(response.choices[0].message.content)
            
        except Exception as e:
            logger.error(f"Error getting census data: {e}")
            return {}

    def generate_comprehensive_content(self, address: str, property_data: PropertyReportData) -> PropertyReportData:
        """
        Generate all content sections using AI with the exact formatting requirements
        """
        logger.info(f"Generating comprehensive content for: {address}")
        
        # Create detailed context for AI
        context = f"""
        Property Address: {address}
        County: {property_data.county}
        State: {property_data.state}
        Coordinates: {property_data.latitude}, {property_data.longitude}
        """
        
        # Generate each section separately to maintain exact formatting
        sections = {
            'property_summary': self._generate_property_summary(context),
            'location_summary': self._generate_location_summary(context),
            'demographic_analysis': self._generate_demographic_analysis(context),
            'size_and_topography': self._generate_size_topography(context),
            'population_analysis': self._generate_population_analysis(context),
            'household_trends': self._generate_household_trends(context),
            'employment_analysis': self._generate_employment_analysis(context),
            'economic_factors': self._generate_economic_factors(context),
            'community_services': self._generate_community_services(context),
            'swot_analysis': self._generate_swot_analysis(context)
        }
        
        # Update property data with generated content
        property_data.property_summary = sections['property_summary']
        property_data.location_summary = sections['location_summary']
        property_data.demographic_analysis = sections['demographic_analysis']
        property_data.size_and_topography = sections['size_and_topography']
        property_data.population_analysis = sections['population_analysis']
        property_data.household_trends = sections['household_trends']
        property_data.employment_analysis = sections['employment_analysis']
        property_data.economic_factors = sections['economic_factors']
        property_data.community_services = sections['community_services']
        
        # SWOT Analysis
        swot = sections['swot_analysis']
        property_data.swot_strengths = swot.get('strengths', '')
        property_data.swot_weaknesses = swot.get('weaknesses', '')
        property_data.swot_opportunities = swot.get('opportunities', '')
        property_data.swot_threats = swot.get('threats', '')
        
        return property_data

    def _generate_property_summary(self, context: str) -> str:
        """Generate property summary matching the exact format"""
        example = """The subject is located in Salem, in Utah County. It is part of the Provo-Orem MSA. The subject property is located in northern Utah within the official boundaries of Utah County. The county is situated directly south of Salt Lake County. This area is generally called the Provo/Orem metropolitan area and is approximately 45 miles south of metropolitan Salt Lake, which is the financial center for the Intermountain Region. This region encompasses all of Utah, southern Idaho, southwestern Wyoming, and eastern Nevada. Utah County is part of a four-county area that is commonly known as the Wasatch Front. Provo is the Utah County seat."""
        
        prompt = f"""
        Generate a property summary paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Maintain the same structure: location, county, MSA/metropolitan area, regional context, broader region description, and county seat information.
        """
        
        return self._get_ai_response(prompt)

    def _generate_location_summary(self, context: str) -> str:
        """Generate location summary matching the exact format"""
        example = """The subject is located on the corner of Hwy 198 and Elk Ridge Drive with good access and exposure. A major thoroughfare in the area is Hwy 198 which partially fronts the subject. The location also offers very close proximity to Salem Pond, Salem High School, Salem Community Center, Salem City Recreation, with limited retail areas in close proximity. The subject is surrounded by vacant land and residential uses. Utah County is broken up into three sectors.  North County (Lindon to Lehi) Central County (Provo/Orem) and South County (Springville to Payson). Central county accounts for a lot of the class B office buildings. The following is taken from Reis, it shows the market area with an arrow pointing to the subject."""
        
        prompt = f"""
        Generate a location summary paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: specific location details, access/exposure, nearby amenities, surrounding land uses, and county/market subdivisions.
        """
        
        return self._get_ai_response(prompt)

    def _generate_demographic_analysis(self, context: str) -> str:
        """Generate demographic analysis - note this matches property_summary in the example"""
        return self._generate_property_summary(context)

    def _generate_size_topography(self, context: str) -> str:
        """Generate size and topography description"""
        example = """The surrounding mountains form a valley about 30 miles wide and 50 miles long. Utah Lake is located centrally to the valley and is Utah's largest freshwater lake. The Wasatch Mountains, which provide a beautiful background to the county on the east, nearly converge with Utah Lake on the west to form the southern boundary south of Santaquin City. The northern boundary is considered the "point of the mountain" which is just north of Lehi City. The elevation varies from 4,480 to 11,928 feet (Mt. Nebo) above sea level. Utah Lake and Mt. Timpanogos present a mountainous scenic backdrop within this metropolitan setting."""
        
        prompt = f"""
        Generate a size and topography paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: geographical features, valley dimensions, major landmarks, elevation ranges, and scenic elements.
        """
        
        return self._get_ai_response(prompt)

    def _generate_population_analysis(self, context: str) -> str:
        """Generate population analysis with specific statistics"""
        example = """According to Pitney Bowes/Gadberry Group - GroundView®, a Geographic Information System (GIS) Company, Utah County had a 2020 total population of 649,258 and experienced an annual growth rate of 2.3%, which was higher than the Utah annual growth rate of 1.6%. The county accounted for 20.0% of the total Utah population (3,254,284). Within the county the population density was 304 people per square mile compared to the lower Utah population density of 38 people per square mile and the lower United States population density of 92 people per square mile."""
        
        prompt = f"""
        Generate a population analysis paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: data source attribution, specific population numbers, growth rates, state comparisons, population density comparisons. Use current 2024 data where available.
        """
        
        return self._get_ai_response(prompt)

    def _generate_household_trends(self, context: str) -> str:
        """Generate household trends analysis"""
        example = """The 2020 number of households in the county was 178,689. The number of households in the county is projected to grow by 2.0% annually, increasing the number of households to 197,669 by 2025. The 2020 average household size for the county was 3.55, which was 37.61% larger than the United States average household size of 2.58 for 2020. The average household size in the county is anticipated to retract by 0.06% annually, reducing the average household size to 3.54 by 2025."""
        
        prompt = f"""
        Generate a household trends paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: specific household numbers, growth projections, average household size, national comparisons, and future projections. Use current 2024 data and project to 2029.
        """
        
        return self._get_ai_response(prompt)

    def _generate_employment_analysis(self, context: str) -> str:
        """Generate employment analysis"""
        example = """Total employment has increased annually over the past decade in the state of Utah by 2.5% and increased annually by 3.9% in the county. From 2018 to 2019 unemployment decreased in Utah by 0.4% and decreased by 0.4% in the county. In the state of Utah unemployment has decreased over the previous month by 1.0% and decreased by 0.7% in the county."""
        
        prompt = f"""
        Generate an employment analysis paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: employment growth rates, unemployment trends, state vs county comparisons, recent changes. Use current 2024 employment data.
        """
        
        return self._get_ai_response(prompt)

    def _generate_economic_factors(self, context: str) -> str:
        """Generate economic factors analysis"""
        example = """Salem is a suburb of Payson and Provo/Orem market area. Salem is still considered somewhat of a rural area but over the years has begun to be built out. A majority of resident's commute to other cities within the metropolitan area for employment. The largest industries in the city include manufacturing, public administration agricultural uses and retail trade. The local economy consists of commercial and industrial businesses on the main arterials. The city's commercial area is on Hwy 198, featuring retail, office, residential, and financial services."""
        
        prompt = f"""
        Generate an economic factors paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: suburban/rural character, development patterns, commuting patterns, major industries, commercial areas, and business types.
        """
        
        return self._get_ai_response(prompt)

    def _generate_community_services(self, context: str) -> str:
        """Generate community services description"""
        example = """Community services and facilities are readily available in the surrounding area. These include public services such as fire stations, hospitals, police stations, and schools (all ages). GreatSchools.org is an on-line tool that rates every school on a scale of one to ten based on test scores. They also track parents rating of the school on a one to five scale."""
        
        prompt = f"""
        Generate a community services paragraph following this EXACT format and style:
        {example}
        
        Use this context: {context}
        
        Include: availability of services, specific service types, educational resources, and rating systems.
        """
        
        return self._get_ai_response(prompt)

    def _generate_swot_analysis(self, context: str) -> Dict[str, str]:
        """Generate SWOT analysis components"""
        prompt = f"""
        Generate a SWOT analysis for this property location. Return as JSON with keys: strengths, weaknesses, opportunities, threats.
        
        Context: {context}
        
        Examples:
        - Strengths: "Easy access from Highway 198 Within close proximity to residential developments"
        - Weaknesses: "The subject has average to weak visibility"  
        - Opportunities: "Opportunity for development of improvement on property"
        - Threats: "There is excess land around the property that could possibility be developed."
        
        Keep each section short and concise, focusing on location-specific factors.
        """
        
        response = self._get_ai_response(prompt, json_response=True)
        try:
            return json.loads(response)
        except:
            return {
                'strengths': 'Good location with development potential',
                'weaknesses': 'Limited visibility from main roads',
                'opportunities': 'Opportunity for future development',
                'threats': 'Competition from nearby available land'
            }

    def _get_ai_response(self, prompt: str, json_response: bool = False) -> str:
        """Get response from OpenAI API"""
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a professional real estate analyst. Maintain exact formatting and style as shown in examples. Provide factual, current information." + (" Respond in valid JSON format." if json_response else "")},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error getting AI response: {e}")
            return "Content generation failed"

    def create_property_report(self, 
                             address: str,
                             prepared_by: str = "",
                             prepared_by_company: str = "",
                             prepared_by_address: str = "",
                             prepared_for: str = "",
                             prepared_for_company: str = "",
                             prepared_for_address: str = "",
                             property_name: str = "",
                             property_type: str = "Vacant Land",
                             **kwargs) -> PropertyReportData:
        """
        Create complete property report data
        """
        logger.info(f"Creating property report for: {address}")
        
        # Get coordinates and location details
        lat, lng, address_details = self.get_coordinates_and_details(address)
        
        # Initialize property data
        property_data = PropertyReportData(
            address=address,
            prepared_by=prepared_by,
            prepared_by_company=prepared_by_company,
            prepared_by_address=prepared_by_address,
            prepared_for=prepared_for,
            prepared_for_company=prepared_for_company,
            prepared_for_address=prepared_for_address,
            property_name=property_name or f"{address_details.get('city', 'Property')} {property_type}",
            property_type=property_type,
            state=address_details.get('state', ''),
            county=address_details.get('county', ''),
            latitude=str(lat),
            longitude=str(lng),
            **kwargs
        )
        
        # Get property images
        aerial_path, street_view_path = self.get_property_images(address, lat, lng)
        property_data.aerial_image_path = aerial_path
        property_data.street_view_image_path = street_view_path
        
        # Generate all content sections
        property_data = self.generate_comprehensive_content(address, property_data)
        
        return property_data

    def _replace_image_placeholder(self, doc: Document, placeholder: str, image_path: Optional[str], width_inches: float = 3.0):
        """
        Replace image placeholder with actual image in the document
        
        Args:
            doc: The Word document
            placeholder: The placeholder text to replace (e.g., '{{aerial_image}}')
            image_path: Path to the image file
            width_inches: Width of the image in inches
        """
        if not image_path or not os.path.exists(image_path):
            logger.warning(f"Image not found for placeholder {placeholder}: {image_path}")
            return
            
        # Search and replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Clear the paragraph
                paragraph.text = paragraph.text.replace(placeholder, '')
                # Add the image
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                logger.info(f"Replaced {placeholder} with image from {image_path}")
                
        # Search and replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            # Clear the paragraph
                            paragraph.text = paragraph.text.replace(placeholder, '')
                            # Add the image
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=Inches(width_inches))
                            logger.info(f"Replaced {placeholder} with image in table from {image_path}")

    def create_word_document(self, property_data: PropertyReportData) -> str:
        """
        Create Word document from property data using template
        """
        logger.info(f"Creating Word document for: {property_data.address}")
        
        # Load template
        doc = Document(self.template_path)
        
        # Define all replacements
        replacements = {
            '{{Date}}': property_data.date,
            '{{prepared_by}}': property_data.prepared_by,
            '{{prepared_by_company}}': property_data.prepared_by_company,
            '{{prepared_by_address}}': property_data.prepared_by_address,
            '{{prepared_for}}': property_data.prepared_for,
            '{{prepared_for_company}}': property_data.prepared_for_company,
            '{{prepared_for_address}}': property_data.prepared_for_address,
            '{{property_summary}}': property_data.property_summary,
            '{{property_name}}': property_data.property_name,
            '{{property_type}}': property_data.property_type,
            '{{state}}': property_data.state,
            '{{county}}': property_data.county,
            '{{longitude}}': property_data.longitude,
            '{{latitude}}': property_data.latitude,
            '{{Topography}}': property_data.topography,
            '{{shape}}': property_data.shape,
            '{{Access}}': property_data.access,
            '{{Exposure}}': property_data.exposure,
            '{{lot_area}}': property_data.lot_area,
            '{{acres}}': property_data.acres,
            '{{recorded_sale_date}}': property_data.recorded_sale_date,
            '{{zoning}}': property_data.zoning,
            '{{apn}}': property_data.apn,
            '{{current_owner}}': property_data.current_owner,
            '{{marketing_period}}': property_data.marketing_period,
            '{{swot_strengths}}': property_data.swot_strengths,
            '{{swot_weaknesses}}': property_data.swot_weaknesses,
            '{{swot_opportunities}}': property_data.swot_opportunities,
            '{{swot_threats}}': property_data.swot_threats,
            '{{location_summary}}': property_data.location_summary,
            '{{demographic_analysis}}': property_data.demographic_analysis,
            '{{size_and_topography}}': property_data.size_and_topography,
            '{{population_analysis}}': property_data.population_analysis,
            '{{household_trends}}': property_data.household_trends,
            '{{employment_analysis}}': property_data.employment_analysis,
            '{{economic_factors}}': property_data.economic_factors,
            '{{community_services}}': property_data.community_services,
        }
        
        # Replace text in all document elements
        self._replace_text_in_document(doc, replacements)
        
        # Replace image placeholders
        self._replace_image_placeholder(doc, '{{ariel_image}}', property_data.aerial_image_path, width_inches=4.0)
        self._replace_image_placeholder(doc, '{{street_view}}', property_data.street_view_image_path, width_inches=4.0)
        
        # Generate output filename
        safe_address = "".join(c for c in property_data.address if c.isalnum() or c in (' ', '-', '_')).rstrip()
        output_filename = f"Property_Report_{safe_address}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / output_filename
        
        # Save document
        doc.save(output_path)
        logger.info(f"Document saved: {output_path}")
        
        return str(output_path)

    def _replace_text_in_document(self, doc: Document, replacements: Dict[str, str]):
        """Replace text in all parts of the document"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)

    def process_single_property(self, address: str, **kwargs) -> str:
        """
        Complete workflow for single property
        """
        property_data = self.create_property_report(address, **kwargs)
        document_path = self.create_word_document(property_data)
        return document_path

    def process_csv_batch(self, csv_path: str) -> List[str]:
        """
        Process multiple properties from CSV file
        CSV should have columns: address, and optionally other property details
        """
        logger.info(f"Processing batch from CSV: {csv_path}")
        
        df = pd.read_csv(csv_path)
        if 'address' not in df.columns:
            raise ValueError("CSV must contain 'address' column")
        
        document_paths = []
        for index, row in df.iterrows():
            try:
                address = row['address']
                logger.info(f"Processing {index + 1}/{len(df)}: {address}")
                
                # Extract other parameters from CSV
                kwargs = {col: str(row[col]) for col in df.columns if col != 'address' and pd.notna(row[col])}
                
                document_path = self.process_single_property(address, **kwargs)
                document_paths.append(document_path)
                
            except Exception as e:
                logger.error(f"Failed to process {address}: {e}")
                continue
        
        return document_paths


def main():
    """
    Example usage
    """
    # Configuration
    OPENAI_API_KEY = 'sk-rSNepa2p0yahWHyJ27CuZd2snzGme9R2hlNOpreUWDT3BlbkFJLVmbTqYzzOpIt-K5mexrZw5W37ziLI2jkNiz0NmmAA'
    GOOGLE_API_KEY = 'AIzaSyCl6Oc03tJ-MkQEXMc84pF9lXURvPLPmHU'
    TEMPLATE_PATH = "template.docx"
    
    if not OPENAI_API_KEY or not GOOGLE_API_KEY:
        raise ValueError("Both OPENAI_API_KEY and GOOGLE_API_KEY environment variables must be set")
    
    # Initialize generator
    generator = ComprehensivePropertyReportGenerator(
        openai_api_key=OPENAI_API_KEY,
        google_api_key=GOOGLE_API_KEY,
        template_path=TEMPLATE_PATH,
        output_dir="property_reports"
    )
    
    # Example usage
    address = "501 N American Fork Ut 84003"
    
    try:
        document_path = generator.process_single_property(
            address=address,
            prepared_by="Brayden Fisher",
            prepared_by_company="Colliers International",
            prepared_by_address="123 North 123 West Orem, UT 12345",
            prepared_for="Austin Shouse",
            prepared_for_company="UCCU Bank",
            prepared_for_address="789 yellow street Provo, UT 12345",
            property_name="Boothes House",
            property_type="Single Family Residantial",
            lot_area="708711",
            acres="16",
            recorded_sale_date="1/24/2011",
            zoning="RA-5",
            apn="30-034-0073",
            current_owner="EKN FAMILY INVESTMENTS LLC"
        )
        
        print(f"Generated comprehensive report: {document_path}")
        
    except Exception as e:
        print(f"Error generating report: {e}")


if __name__ == "__main__":
    main()