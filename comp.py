import os
import re
import logging
from pathlib import Path
from typing import List, Optional, Dict
from dataclasses import dataclass
from datetime import datetime

import pdfplumber
import fitz  # PyMuPDF for image extraction
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class CompProperty:
    """Data structure for comparable property"""
    comp_number: int
    property_name: str = ""
    primary_use: str = ""
    address: str = ""
    market: str = ""
    sub_market: str = ""
    comp_sf: str = ""
    acres: str = ""
    sale_price: str = ""
    sale_price_sf: str = ""
    sale_price_acres: str = ""
    zoning: str = ""
    parcel_number: str = ""
    off_market_date: str = ""
    months_on_market: str = ""
    topography: str = ""
    seller_landlord: str = ""
    buyer_tenant: str = ""
    listing_broker: str = ""
    listing_broker_company: str = ""
    listing_broker_phone: str = ""
    listing_broker_email: str = ""
    image_path: Optional[str] = None

class CompExtractor:
    """Extract comparable property data from PDF files and replace keywords in Word documents"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.images_dir = self.output_dir / "comp_images"
        self.images_dir.mkdir(exist_ok=True, parents=True)
        
    def extract_comps_from_pdf(self, pdf_path: str) -> List[CompProperty]:
        """
        Extract comparable property data from a PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of CompProperty objects
        """
        logger.info(f"Extracting comp data from PDF: {pdf_path}")
        comps = []
        
        try:
            # Extract text using pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                full_text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n"
                
                # Parse all comps from the full text
                comps = self._parse_all_comps(full_text)
            
            # Extract images using PyMuPDF
            self._extract_images_from_pdf(pdf_path, comps)
            
            logger.info(f"Successfully extracted {len(comps)} comps from PDF")
            
        except Exception as e:
            logger.error(f"Error extracting data from PDF: {e}")
            
        return comps
    
    def _parse_all_comps(self, text: str) -> List[CompProperty]:
        """Parse all comps from the PDF text"""
        comps = []
        
        # Find all comp entries using pattern
        # Pattern looks for "1. Sold Land Property" or "1. Sold Land Space"
        comp_pattern = r'(\d+)\.\s+Sold Land (?:Property|Space)[^\n]*\n(.*?)(?=\d+\.\s+Sold Land|$)'
        
        matches = list(re.finditer(comp_pattern, text, re.DOTALL))
        
        for match in matches:
            comp_num = int(match.group(1))
            comp_text = match.group(0)
            
            comp = self._parse_single_comp(comp_text, comp_num)
            if comp:
                comps.append(comp)
                logger.info(f"Parsed comp {comp_num}: {comp.property_name} at {comp.address}")
        
        return comps
    
    def _parse_single_comp(self, text: str, comp_num: int) -> Optional[CompProperty]:
        """Parse a single comp from text block"""
        comp = CompProperty(comp_number=comp_num)
        
        try:
            # Extract property name/description
            name_match = re.search(r'System ID: \d+\)\s*(.+?)(?:\n|$)', text)
            if name_match:
                comp.property_name = name_match.group(1).strip()
            
            # Extract primary use
            use_match = re.search(r'Primary Use:\s*(.+?)(?:\n|$)', text)
            if use_match:
                comp.primary_use = use_match.group(1).strip()
            
            # Extract address - improved pattern
            # Look for lines between the property name and "Market:" that contain city/state info
            lines = text.split('\n')
            address_parts = []
            capture_address = False
            
            for i, line in enumerate(lines):
                line = line.strip()
                
                # Start capturing after primary use or property name
                if 'Primary Use:' in line or (name_match and name_match.group(1) in line):
                    capture_address = True
                    continue
                
                # Stop capturing when we hit Market:
                if 'Market:' in line:
                    break
                
                # Capture address lines
                if capture_address and line:
                    # Skip if it's a field label
                    if not any(label in line for label in ['Comp SF:', 'Acres:', 'Sale Price:', 'Zoning:', 'Parcel #:']):
                        # Check if line contains state abbreviation (indicates it's likely an address)
                        if re.search(r',\s*[A-Z]{2}\s+\d{5}', line) or len(address_parts) < 2:
                            address_parts.append(line)
            
            comp.address = ' '.join(address_parts)
            
            # Extract market and submarket
            market_match = re.search(r'Market:\s*([^/\n]+?)(?:/\s*Sub-Market:\s*(.+?))?(?:\n|$)', text)
            if market_match:
                comp.market = market_match.group(1).strip()
                if market_match.group(2):
                    comp.sub_market = market_match.group(2).strip()
            
            # Extract all the numeric and property details
            comp.comp_sf = self._extract_field(text, r'Comp SF:\s*(.+?)(?:\s*SF)?(?:\n|$)')
            comp.acres = self._extract_field(text, r'Acres:\s*(.+?)(?:\s*Acres)?(?:\n|$)')
            comp.sale_price = self._extract_field(text, r'Sale Price:\s*(.+?)(?:\n|$)')
            comp.sale_price_sf = self._extract_field(text, r'Sale Price/SF:\s*(.+?)(?:\n|$)')
            comp.sale_price_acres = self._extract_field(text, r'Sale Price/Acres:\s*(.+?)(?:\n|$)')
            comp.zoning = self._extract_field(text, r'Zoning:\s*(.+?)(?:\n|$)')
            comp.parcel_number = self._extract_field(text, r'Parcel #:\s*(.+?)(?:\n|$)')
            comp.off_market_date = self._extract_field(text, r'Off-Market:\s*(.+?)(?:\n|$)')
            comp.months_on_market = self._extract_field(text, r'Months on Market:\s*(.+?)(?:\n|$)')
            comp.topography = self._extract_field(text, r'Topography:\s*(.+?)(?:\n|$)')
            
            # Extract parties
            comp.seller_landlord = self._extract_party(text, 'Seller/Landlord')
            comp.buyer_tenant = self._extract_party(text, 'Buyer/Tenant')
            
            # Extract listing broker info
            broker_info = self._extract_broker_info(text)
            if broker_info:
                comp.listing_broker = broker_info.get('name', '')
                comp.listing_broker_company = broker_info.get('company', '')
                comp.listing_broker_phone = broker_info.get('phone', '')
                comp.listing_broker_email = broker_info.get('email', '')
            
            return comp
            
        except Exception as e:
            logger.error(f"Error parsing comp {comp_num}: {e}")
            return None
    
    def _extract_field(self, text: str, pattern: str) -> str:
        """Extract a field value using regex"""
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            # Remove any curly braces if present
            value = value.replace('{', '').replace('}', '')
            return value
        return ""
    
    def _extract_party(self, text: str, party_type: str) -> str:
        """Extract party (seller/buyer) information"""
        # Look for the party type followed by company name
        pattern = rf'{party_type}\s+([^\n]+?)(?:\n|$)'
        match = re.search(pattern, text, re.MULTILINE)
        if match:
            # The company name is usually on the same line or next line
            company_line = match.group(1).strip()
            if company_line and not any(x in company_line.lower() for x in ['role', 'company', 'name', 'phone', 'email']):
                return company_line
            
            # If not found, check the next line
            lines_after = text[match.end():].split('\n')
            for line in lines_after[:2]:
                line = line.strip()
                if line and not any(x in line.lower() for x in ['role', 'company', 'name', 'phone', 'email', 'listing', 'procuring']):
                    return line
        return ""
    
    def _extract_broker_info(self, text: str) -> Dict[str, str]:
        """Extract listing broker information"""
        info = {}
        
        # Find listing broker section
        broker_match = re.search(r'Listing Broker\s+([^\n]+)', text, re.MULTILINE)
        if broker_match:
            # Extract company (usually right after "Listing Broker")
            info['company'] = broker_match.group(1).strip()
            
            # Look for broker details in subsequent lines
            text_after = text[broker_match.end():]
            lines = text_after.split('\n')[:5]  # Check next 5 lines
            
            for line in lines:
                line = line.strip()
                # Email
                email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', line)
                if email_match:
                    info['email'] = email_match.group(0)
                # Phone
                phone_match = re.search(r'(\d{3}[.-]?\d{3}[.-]?\d{4})', line)
                if phone_match:
                    info['phone'] = phone_match.group(1)
                # Name (if it's not a label and not already captured)
                elif line and not any(x in line.lower() for x in ['role', 'company', 'phone', 'email', 'broker']) and 'name' not in info:
                    # Check if it looks like a person's name (has space, starts with capital)
                    if ' ' in line and line[0].isupper():
                        info['name'] = line
        
        return info
    
    def _extract_images_from_pdf(self, pdf_path: str, comps: List[CompProperty]):
        """Extract images from PDF and associate with comps"""
        try:
            pdf_document = fitz.open(pdf_path)
            
            # Track which comp we're on based on page
            comp_index = 0
            images_per_comp = []
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                image_list = page.get_images(full=True)
                
                if image_list:
                    # Assume each page with images corresponds to a comp
                    # (you may need to adjust this logic based on your PDF structure)
                    for img_index, img in enumerate(image_list):
                        if comp_index < len(comps):
                            try:
                                # Extract the image
                                xref = img[0]
                                base_image = pdf_document.extract_image(xref)
                                image_bytes = base_image["image"]
                                image_ext = base_image["ext"]
                                
                                # Save the image
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                image_filename = f"comp_{comps[comp_index].comp_number}_{timestamp}.{image_ext}"
                                image_path = self.images_dir / image_filename
                                
                                with open(image_path, "wb") as img_file:
                                    img_file.write(image_bytes)
                                
                                # Associate with comp (only first image per comp)
                                if not comps[comp_index].image_path:
                                    comps[comp_index].image_path = str(image_path)
                                    logger.info(f"Extracted image for comp {comps[comp_index].comp_number}")
                                
                            except Exception as e:
                                logger.error(f"Error extracting image {img_index} on page {page_num}: {e}")
                    
                    # Move to next comp after processing images on this page
                    comp_index += 1
            
            pdf_document.close()
            
        except Exception as e:
            logger.error(f"Error extracting images from PDF: {e}")
    
    def create_comp_replacements(self, comps: List[CompProperty]) -> Dict[str, str]:
        """
        Create a dictionary of keyword replacements from comp data
        
        Args:
            comps: List of CompProperty objects
            
        Returns:
            Dictionary mapping keywords to values
        """
        replacements = {}
        
        # Limit to 6 comps as requested
        for i, comp in enumerate(comps[:6], 1):
            # Create all the replacement keywords for this comp
            replacements[f'{{{{comp{i}_number}}}}'] = str(comp.comp_number)
            replacements[f'{{{{comp{i}_property_name}}}}'] = comp.property_name
            replacements[f'{{{{comp{i}_address}}}}'] = comp.address
            replacements[f'{{{{comp{i}_primary_use}}}}'] = comp.primary_use
            replacements[f'{{{{comp{i}_market}}}}'] = comp.market
            replacements[f'{{{{comp{i}_sub_market}}}}'] = comp.sub_market
            replacements[f'{{{{comp{i}_comp_sf}}}}'] = comp.comp_sf
            replacements[f'{{{{comp{i}_acres}}}}'] = comp.acres
            replacements[f'{{{{comp{i}_sale_price}}}}'] = comp.sale_price
            replacements[f'{{{{comp{i}_sale_price_sf}}}}'] = comp.sale_price_sf
            replacements[f'{{{{comp{i}_sale_price_acres}}}}'] = comp.sale_price_acres
            replacements[f'{{{{comp{i}_zoning}}}}'] = comp.zoning
            replacements[f'{{{{comp{i}_parcel_number}}}}'] = comp.parcel_number
            replacements[f'{{{{comp{i}_off_market_date}}}}'] = comp.off_market_date
            replacements[f'{{{{comp{i}_months_on_market}}}}'] = comp.months_on_market
            replacements[f'{{{{comp{i}_topography}}}}'] = comp.topography
            replacements[f'{{{{comp{i}_seller_landlord}}}}'] = comp.seller_landlord
            replacements[f'{{{{comp{i}_buyer_tenant}}}}'] = comp.buyer_tenant
            replacements[f'{{{{comp{i}_listing_broker}}}}'] = comp.listing_broker
            replacements[f'{{{{comp{i}_listing_broker_company}}}}'] = comp.listing_broker_company
            replacements[f'{{{{comp{i}_listing_broker_phone}}}}'] = comp.listing_broker_phone
            replacements[f'{{{{comp{i}_listing_broker_email}}}}'] = comp.listing_broker_email
            
            # Add image path if available
            if comp.image_path:
                replacements[f'{{{{comp{i}_image}}}}'] = comp.image_path
        
        return replacements
    
    def replace_keywords_in_document(self, doc_path: str, comps: List[CompProperty], output_path: Optional[str] = None) -> str:
        """
        Replace keywords in a Word document with comp data
        
        Args:
            doc_path: Path to the Word document template
            comps: List of CompProperty objects
            output_path: Optional output path for the modified document
            
        Returns:
            Path to the saved document
        """
        logger.info(f"Replacing keywords in document: {doc_path}")
        
        # Load the document
        doc = Document(doc_path)
        
        # Create replacements dictionary
        replacements = self.create_comp_replacements(comps)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for keyword, value in replacements.items():
                if keyword in paragraph.text:
                    # Replace the keyword with the value
                    inline = paragraph.runs
                    for run in inline:
                        if keyword in run.text:
                            run.text = run.text.replace(keyword, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for keyword, value in replacements.items():
                            if keyword in paragraph.text:
                                inline = paragraph.runs
                                for run in inline:
                                    if keyword in run.text:
                                        run.text = run.text.replace(keyword, value)
        
        # Handle image replacements
        for i in range(1, 7):  # For comps 1-6
            image_keyword = f'{{{{comp{i}_image}}}}'
            image_path = replacements.get(image_keyword, '')
            
            if image_path and os.path.exists(image_path):
                # Replace image placeholder with actual image
                for paragraph in doc.paragraphs:
                    if image_keyword in paragraph.text:
                        # Clear the paragraph text
                        paragraph.text = paragraph.text.replace(image_keyword, '')
                        # Add the image
                        try:
                            paragraph.add_run().add_picture(image_path, width=Inches(4.0))
                            logger.info(f"Replaced {image_keyword} with image")
                        except Exception as e:
                            logger.error(f"Error adding image for {image_keyword}: {e}")
        
        # Save the document
        if not output_path:
            output_path = self.output_dir / f"Report_with_Comps_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc.save(output_path)
        logger.info(f"Document saved with replaced keywords: {output_path}")
        
        # Log which replacements were made
        logger.info("Replacements made:")
        for keyword, value in replacements.items():
            if value:  # Only log non-empty values
                logger.info(f"  {keyword} -> {value}")
        
        return str(output_path)

def main():
    """Example usage"""
    # Initialize the extractor
    extractor = CompExtractor(output_dir="property_reports")
    
    # Extract comps from PDF
    pdf_path = "comp.pdf"  # Your PDF file
    comps = extractor.extract_comps_from_pdf(pdf_path)
    
    # Print extracted data
    print(f"\nExtracted {len(comps)} comparable properties:")
    for comp in comps[:6]:  # Show first 6
        print(f"\nComp {comp.comp_number}: {comp.property_name}")
        print(f"Address: {comp.address}")
        print(f"Sale Price: {comp.sale_price}")
        print(f"Price/SF: {comp.sale_price_sf}")
    
    # Replace keywords in your Word template
    template_path = "comptemplate.docx"  # Your Word template with {{comp1_address}} style keywords
    output_path = extractor.replace_keywords_in_document(template_path, comps)
    print(f"\nDocument created with replaced keywords: {output_path}")

if __name__ == "__main__":
    main()