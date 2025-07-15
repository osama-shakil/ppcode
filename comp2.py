import os
import re
import logging
from pathlib import Path
from typing import List, Optional, Dict, Tuple
from dataclasses import dataclass
from datetime import datetime

import pdfplumber
import fitz  # PyMuPDF for image extraction
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class CompProperty:
    """Data structure for comparable property"""
    comp_number: int
    property_name: str = ""
    primary_use: str = ""
    address: str = ""
    market: str = ""
    sub_market: str = ""
    comp_sf: str = ""
    acres: str = ""
    sale_price: str = ""
    sale_price_sf: str = ""
    sale_price_acres: str = ""
    zoning: str = ""
    parcel_number: str = ""
    off_market_date: str = ""
    months_on_market: str = ""
    topography: str = ""
    seller_landlord: str = ""
    buyer_tenant: str = ""
    listing_broker: str = ""
    listing_broker_company: str = ""
    listing_broker_phone: str = ""
    listing_broker_email: str = ""
    image_path: Optional[str] = None

class CompExtractor:
    """Extract comparable property data from PDF files and replace keywords in Word documents"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.images_dir = self.output_dir / "comp_images"
        self.images_dir.mkdir(exist_ok=True, parents=True)
        
    def extract_comps_from_pdf(self, pdf_path: str) -> List[CompProperty]:
        """
        Extract comparable property data from a PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of CompProperty objects
        """
        logger.info(f"Extracting comp data from PDF: {pdf_path}")
        comps = []
        
        try:
            # Extract text using pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                full_text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n"
                
                # Clean up common PDF artifacts
                full_text = self._clean_pdf_text(full_text)
                
                # Parse all comps from the full text
                comps = self._parse_all_comps(full_text)
            
            # Extract images using PyMuPDF
            self._extract_images_from_pdf(pdf_path, comps)
            
            logger.info(f"Successfully extracted {len(comps)} comps from PDF")
            
        except Exception as e:
            logger.error(f"Error extracting data from PDF: {e}")
            
        return comps
    
    def _clean_pdf_text(self, text: str) -> str:
        """Remove common PDF artifacts and formatting issues"""
        # Remove page footers and headers
        text = re.sub(r'All information contained herein.*?Page \d+ of \d+', '', text, flags=re.DOTALL)
        text = re.sub(r'Land Comp Summary Report', '', text)
        text = re.sub(r'July \d+, \d{4}', '', text)
        text = re.sub(r'Colliers\s*\n', '', text)
        
        return text
    
    def _parse_all_comps(self, text: str) -> List[CompProperty]:
        """Parse all comps from the PDF text"""
        comps = []
        
        # Find all comp entries - updated pattern to be more flexible
        comp_pattern = r'(\d+)\.\s+Sold Land (?:Property|Space)(?:\s*\([^)]+\))?\s*([^\n]*)'
        
        # Split text into comp sections
        comp_sections = []
        matches = list(re.finditer(comp_pattern, text, re.MULTILINE))
        
        for i, match in enumerate(matches):
            start = match.start()
            # End is either the start of next comp or end of text
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            
            comp_num = int(match.group(1))
            comp_text = text[start:end]
            comp_sections.append((comp_num, comp_text))
        
        # Parse each comp section
        for comp_num, comp_text in comp_sections:
            comp = self._parse_single_comp(comp_text, comp_num)
            if comp:
                comps.append(comp)
                logger.info(f"Parsed comp {comp_num}: {comp.property_name} at {comp.address}")
        
        return comps
    
    def _parse_single_comp(self, text: str, comp_num: int) -> Optional[CompProperty]:
        """Parse a single comp from text block"""
        comp = CompProperty(comp_number=comp_num)
        
        try:
            # Extract property name - it's usually after the System ID
            name_match = re.search(r'System ID:\s*\d+\)\s*([^\n]+)', text)
            if name_match:
                comp.property_name = name_match.group(1).strip()
            
            # Extract primary use
            use_match = re.search(r'Primary Use:\s*([^\n]+)', text)
            if use_match:
                comp.primary_use = use_match.group(1).strip()
            
            # Extract address - improved method
            comp.address = self._extract_address(text)
            
            # Extract market and submarket
            market_match = re.search(r'Market:\s*([^/\n]+?)(?:\s*/\s*Sub-Market:\s*([^\n]+))?', text)
            if market_match:
                comp.market = market_match.group(1).strip()
                if market_match.group(2):
                    comp.sub_market = market_match.group(2).strip()
            
            # Extract all the numeric and property details with improved patterns
            comp.comp_sf = self._extract_clean_field(text, r'Comp SF:\s*([0-9,]+(?:\s*SF)?)')
            comp.acres = self._extract_clean_field(text, r'Acres:\s*([0-9.]+(?:\s*Acres)?)')
            comp.sale_price = self._extract_clean_field(text, r'Sale Price:\s*(\$[0-9,]+)')
            comp.sale_price_sf = self._extract_clean_field(text, r'Sale Price/SF:\s*(\$[0-9,.]+)')
            comp.sale_price_acres = self._extract_clean_field(text, r'Sale Price/Acres:\s*(\$[0-9,.]+)')
            comp.zoning = self._extract_clean_field(text, r'Zoning:\s*([^\n]+?)(?:\s*Off-Market:|$)')
            comp.parcel_number = self._extract_clean_field(text, r'Parcel #:\s*([^\n]+?)(?:\s*Lot Dimensions:|$)')
            comp.off_market_date = self._extract_clean_field(text, r'Off-Market:\s*(\d{2}/\d{2}/\d{4})')
            comp.months_on_market = self._extract_clean_field(text, r'Months on Market:\s*(\d+)')
            comp.topography = self._extract_clean_field(text, r'Topography:\s*([^\n]+?)(?:\s*Land Conditions:|$)')
            
            # Extract parties with improved method
            comp.seller_landlord = self._extract_party(text, 'Seller/Landlord')
            comp.buyer_tenant = self._extract_party(text, 'Buyer/Tenant')
            
            # Extract listing broker info with improved method
            broker_info = self._extract_broker_info(text)
            if broker_info:
                comp.listing_broker = broker_info.get('name', '')
                comp.listing_broker_company = broker_info.get('company', '')
                comp.listing_broker_phone = broker_info.get('phone', '')
                comp.listing_broker_email = broker_info.get('email', '')
            
            return comp
            
        except Exception as e:
            logger.error(f"Error parsing comp {comp_num}: {e}")
            return None
    
    def _extract_address(self, text: str) -> str:
        """Extract address with improved logic"""
        # Method 1: Look for a complete address pattern with city, state, zip
        full_address_match = re.search(r'(\d+[^,\n]+,\s*[^,\n]+,\s*[A-Z]{2}\s+\d{5})', text)
        if full_address_match:
            return full_address_match.group(1).strip()
        
        # Method 2: Extract lines between property name/primary use and Market
        lines = text.split('\n')
        address_lines = []
        in_address_section = False
        
        for line in lines:
            line = line.strip()
            
            # Start capturing after Primary Use
            if 'Primary Use:' in line:
                in_address_section = True
                continue
            
            # Stop at Market
            if 'Market:' in line:
                break
            
            # Capture address lines
            if in_address_section and line:
                # Skip if it's a field label
                if not any(label in line for label in ['Comp SF:', 'Acres:', 'Sale Price:', 'Zoning:', 
                                                       'Parcel #:', 'Off-Market:', 'Lot Dimensions:',
                                                       'Legal Description:', 'Gas:', 'Water:', 'Sewer:',
                                                       'Power:', 'Rail Status:', 'Topography:']):
                    address_lines.append(line)
        
        # Join address lines, but limit to first 2-3 lines that look like address
        final_address_parts = []
        for line in address_lines[:3]:
            # Check if line looks like part of an address
            if (re.search(r'\d+', line) or  # Contains numbers
                re.search(r'\b(Street|St|Avenue|Ave|Road|Rd|Drive|Dr|Lane|Ln|Way|Blvd|Boulevard)\b', line, re.I) or
                re.search(r',\s*[A-Z]{2}\s+\d{5}', line)):  # City, State ZIP
                final_address_parts.append(line)
        
        return ' '.join(final_address_parts).strip()
    
    def _extract_clean_field(self, text: str, pattern: str) -> str:
        """Extract a field value and clean it"""
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            # Remove any curly braces
            value = value.replace('{', '').replace('}', '')
            # Remove duplicate colons and clean up
            value = re.sub(r':+', '', value)
            # Remove "SF" or "Acres" if it's redundant
            value = re.sub(r'\s*SF\s*$', '', value)
            value = re.sub(r'\s*Acres\s*$', '', value)
            return value.strip()
        return ""
    
    def _extract_party(self, text: str, party_type: str) -> str:
        """Extract party (seller/buyer) information with improved logic"""
        # Look for the party type in a table row structure
        pattern = rf'{party_type}\s*\n\s*([^\n]+?)(?:\n|$)'
        match = re.search(pattern, text, re.MULTILINE)
        
        if match:
            party_name = match.group(1).strip()
            # Clean up any Role/Company/Name labels
            if party_name and not any(x in party_name.lower() for x in ['role', 'company', 'name', 'phone', 'email']):
                return party_name
        
        # Alternative: Look for party in table structure
        alt_pattern = rf'{party_type}\s+([^\n]+?)(?:\s+Role|\s+Company|\s+Listing|\s+Procuring|\n|$)'
        alt_match = re.search(alt_pattern, text)
        if alt_match:
            return alt_match.group(1).strip()
        
        return ""
    
    def _extract_broker_info(self, text: str) -> Dict[str, str]:
        """Extract listing broker information with improved parsing"""
        info = {}
        
        # Look for the listing broker section
        broker_section_match = re.search(r'Listing Broker\s+(.*?)(?:Procuring Broker|Buyer/Tenant|$)', text, re.DOTALL)
        if not broker_section_match:
            return info
        
        broker_text = broker_section_match.group(1)
        lines = broker_text.split('\n')
        
        # First line after "Listing Broker" is usually the company
        if lines and lines[0].strip():
            company = lines[0].strip()
            if not any(x in company.lower() for x in ['role', 'company', 'name', 'phone', 'email']):
                info['company'] = company
        
        # Look for broker name, phone, and email in subsequent lines
        for i, line in enumerate(lines[1:], 1):
            line = line.strip()
            
            # Email
            email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', line)
            if email_match and 'email' not in info:
                info['email'] = email_match.group(0)
            
            # Phone
            phone_match = re.search(r'(\d{3}[.-]?\d{3}[.-]?\d{4})', line)
            if phone_match and 'phone' not in info:
                info['phone'] = phone_match.group(1)
            
            # Name - usually appears before phone/email
            if (not any(x in line.lower() for x in ['role', 'company', 'phone', 'email', 'listing', 'procuring']) 
                and 'name' not in info 
                and line 
                and not re.match(r'^[\d.-]+$', line)  # Not just numbers
                and '@' not in line):  # Not email
                # Check if it looks like a person's name
                if ' ' in line or (line[0].isupper() and len(line) > 2):
                    info['name'] = line
        
        return info
    
    def _extract_images_from_pdf(self, pdf_path: str, comps: List[CompProperty]):
        """Extract images from PDF and associate with comps"""
        try:
            pdf_document = fitz.open(pdf_path)
            
            # Create a mapping of page numbers to comp numbers
            # This assumes comps appear in order in the PDF
            comp_page_map = {}
            current_comp_idx = 0
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                # Check if this page contains a new comp
                if re.search(r'\d+\.\s+Sold Land (?:Property|Space)', page_text):
                    comp_page_map[page_num] = current_comp_idx
                    current_comp_idx += 1
                elif current_comp_idx > 0:
                    # Continue with the same comp if no new comp found
                    comp_page_map[page_num] = current_comp_idx - 1
            
            # Extract images
            for page_num in range(pdf_document.page_count):
                if page_num not in comp_page_map:
                    continue
                    
                comp_idx = comp_page_map[page_num]
                if comp_idx >= len(comps):
                    continue
                
                page = pdf_document[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    if not comps[comp_idx].image_path:  # Only take first image per comp
                        try:
                            # Extract the image
                            xref = img[0]
                            base_image = pdf_document.extract_image(xref)
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]
                            
                            # Save the image
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            image_filename = f"comp_{comps[comp_idx].comp_number}_{timestamp}.{image_ext}"
                            image_path = self.images_dir / image_filename
                            
                            with open(image_path, "wb") as img_file:
                                img_file.write(image_bytes)
                            
                            comps[comp_idx].image_path = str(image_path)
                            logger.info(f"Extracted image for comp {comps[comp_idx].comp_number}")
                            
                        except Exception as e:
                            logger.error(f"Error extracting image {img_index} on page {page_num}: {e}")
            
            pdf_document.close()
            
        except Exception as e:
            logger.error(f"Error extracting images from PDF: {e}")
    
    def create_comp_replacements(self, comps: List[CompProperty]) -> Dict[str, str]:
        """
        Create a dictionary of keyword replacements from comp data
        
        Args:
            comps: List of CompProperty objects
            
        Returns:
            Dictionary mapping keywords to values
        """
        replacements = {}
        
        # Create a sorted list of first 6 comps based on comp_number
        sorted_comps = sorted(comps, key=lambda x: x.comp_number)[:6]
        
        # Create replacements for each comp position (1-6)
        for i in range(1, 7):
            if i <= len(sorted_comps):
                comp = sorted_comps[i-1]
                
                # Create all the replacement keywords for this comp
                replacements[f'{{{{comp{i}_number}}}}'] = str(comp.comp_number)
                replacements[f'{{{{comp{i}_property_name}}}}'] = comp.property_name
                replacements[f'{{{{comp{i}_address}}}}'] = comp.address
                replacements[f'{{{{comp{i}_primary_use}}}}'] = comp.primary_use
                replacements[f'{{{{comp{i}_market}}}}'] = comp.market
                replacements[f'{{{{comp{i}_sub_market}}}}'] = comp.sub_market
                replacements[f'{{{{comp{i}_comp_sf}}}}'] = comp.comp_sf
                replacements[f'{{{{comp{i}_acres}}}}'] = comp.acres
                replacements[f'{{{{comp{i}_sale_price}}}}'] = comp.sale_price
                replacements[f'{{{{comp{i}_sale_price_sf}}}}'] = comp.sale_price_sf
                replacements[f'{{{{comp{i}_sale_price_acres}}}}'] = comp.sale_price_acres
                replacements[f'{{{{comp{i}_zoning}}}}'] = comp.zoning
                replacements[f'{{{{comp{i}_parcel_number}}}}'] = comp.parcel_number
                replacements[f'{{{{comp{i}_off_market_date}}}}'] = comp.off_market_date
                replacements[f'{{{{comp{i}_months_on_market}}}}'] = comp.months_on_market
                replacements[f'{{{{comp{i}_topography}}}}'] = comp.topography
                replacements[f'{{{{comp{i}_seller_landlord}}}}'] = comp.seller_landlord
                replacements[f'{{{{comp{i}_buyer_tenant}}}}'] = comp.buyer_tenant
                replacements[f'{{{{comp{i}_listing_broker}}}}'] = comp.listing_broker
                replacements[f'{{{{comp{i}_listing_broker_company}}}}'] = comp.listing_broker_company
                replacements[f'{{{{comp{i}_listing_broker_phone}}}}'] = comp.listing_broker_phone
                replacements[f'{{{{comp{i}_listing_broker_email}}}}'] = comp.listing_broker_email
                
                # Add image path if available
                if comp.image_path:
                    replacements[f'{{{{comp{i}_image}}}}'] = comp.image_path
            else:
                # If we don't have enough comps, leave the placeholders empty
                # This prevents partial replacements
                fields = ['number', 'property_name', 'address', 'primary_use', 'market', 'sub_market',
                         'comp_sf', 'acres', 'sale_price', 'sale_price_sf', 'sale_price_acres',
                         'zoning', 'parcel_number', 'off_market_date', 'months_on_market', 
                         'topography', 'seller_landlord', 'buyer_tenant', 'listing_broker',
                         'listing_broker_company', 'listing_broker_phone', 'listing_broker_email']
                
                for field in fields:
                    replacements[f'{{{{comp{i}_{field}}}}}'] = ''
        
        return replacements
    
    def replace_keywords_in_document(self, doc_path: str, comps: List[CompProperty], output_path: Optional[str] = None) -> str:
        """
        Replace keywords in a Word document with comp data
        
        Args:
            doc_path: Path to the Word document template
            comps: List of CompProperty objects
            output_path: Optional output path for the modified document
            
        Returns:
            Path to the saved document
        """
        logger.info(f"Replacing keywords in document: {doc_path}")
        
        # Load the document
        doc = Document(doc_path)
        
        # Create replacements dictionary
        replacements = self.create_comp_replacements(comps)
        
        # Function to replace text in a paragraph while preserving formatting
        def replace_in_paragraph(paragraph):
            """Replace keywords in a paragraph while trying to preserve formatting"""
            if not paragraph.text:
                return
                
            # Check if paragraph contains any keywords
            full_text = paragraph.text
            has_replacement = False
            
            for keyword, value in replacements.items():
                if keyword in full_text:
                    has_replacement = True
                    full_text = full_text.replace(keyword, value)
            
            if has_replacement:
                # Try to preserve formatting by working with runs
                if len(paragraph.runs) == 1:
                    # Simple case: single run
                    paragraph.runs[0].text = full_text
                else:
                    # Complex case: multiple runs
                    # Store the first run's format
                    if paragraph.runs:
                        first_run_format = paragraph.runs[0]
                        # Clear all runs
                        for run in paragraph.runs:
                            run.text = ""
                        # Set the full text in the first run
                        first_run_format.text = full_text
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        
        # Handle image replacements
        for i in range(1, 7):  # For comps 1-6
            image_keyword = f'{{{{comp{i}_image}}}}'
            image_path = replacements.get(image_keyword, '')
            
            if image_path and os.path.exists(image_path):
                # Replace image placeholder with actual image
                for paragraph in doc.paragraphs:
                    if image_keyword in paragraph.text:
                        # Clear the paragraph text
                        paragraph.text = paragraph.text.replace(image_keyword, '')
                        # Add the image
                        try:
                            paragraph.add_run().add_picture(image_path, width=Inches(4.0))
                            logger.info(f"Replaced {image_keyword} with image")
                        except Exception as e:
                            logger.error(f"Error adding image for {image_keyword}: {e}")
        
        # Save the document
        if not output_path:
            output_path = self.output_dir / f"Report_with_Comps_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc.save(output_path)
        logger.info(f"Document saved with replaced keywords: {output_path}")
        
        # Log summary of replacements
        logger.info("Replacement summary:")
        non_empty_replacements = {k: v for k, v in replacements.items() if v}
        logger.info(f"Total replacements made: {len(non_empty_replacements)}")
        
        return str(output_path)

def main():
    """Example usage"""
    # Initialize the extractor
    extractor = CompExtractor(output_dir="property_reports")
    
    # Extract comps from PDF
    pdf_path = "comp.pdf"  # Your PDF file
    comps = extractor.extract_comps_from_pdf(pdf_path)
    
    # Print extracted data
    print(f"\nExtracted {len(comps)} comparable properties:")
    for comp in comps[:6]:  # Show first 6
        print(f"\nComp {comp.comp_number}: {comp.property_name}")
        print(f"Address: {comp.address}")
        print(f"Sale Price: {comp.sale_price}")
        print(f"Price/SF: {comp.sale_price_sf}")
        print(f"Market: {comp.market}")
        if comp.sub_market:
            print(f"Sub-Market: {comp.sub_market}")
    
    # Replace keywords in your Word template
    template_path = "comptemplate.docx"  # Your Word template with {{comp1_address}} style keywords
    output_path = extractor.replace_keywords_in_document(template_path, comps)
    print(f"\nDocument created with replaced keywords: {output_path}")

if __name__ == "__main__":
    main()